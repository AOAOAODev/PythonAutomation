import requests # used for the api 
import random # sort random parks 
from  pprint import pprint
import docx # create doc file 
from PIL import Image
from io import BytesIO

park_guide = docx.Document() # create doc 
park_guide.add_heading('National Park Travel Guide' ,0) # heading 


api_url = "https://national-parks-1150.azurewebsites.net/api/list" #api url 


url_response = requests.get(api_url).json() # make a request to the api and parse it into json


#print(url_response) #print json 
 
five_random_parks = random.sample(url_response, 5) # using random to pick 5 random parks 


#print(five_random_parks) # print random pranks 


for park in five_random_parks: # make a for loop 
    
  base_park_detail_url = 'https://national-parks-1150.azurewebsites.net/api/' # api url 
  park_code = park['park_code'] # park code from the apoi
 #print(park_code) # print park code from api 
  park_name = park['name'] # print name of park
  park_guide.add_paragraph(park_name, 'Heading 1')


  park_detail_url = base_park_detail_url + park_code #deatils we need 
  #print(park_detail_url) # printing again 


  park_detail_response = requests.get(park_detail_url).json() # parse into json and get response 
  #pprint(park_detail_response) # print it out 

  # getting the descriptions from api and putting it into doc 
  park_description = park_detail_response['description']
  park_guide.add_paragraph('Description', 'Heading 1')
  park_guide.add_paragraph(park_description)

  # pretty much the same as everything else but now we add the weather stuff from the api
  park_weather = park_detail_response['weather_overview']
  park_guide.add_paragraph('Weather:', 'Heading 1')
  park_guide.add_paragraph(park_weather)


    # contact info - i look into the api and found the key for it now we just put it on the doc 
  contact_info = park_detail_response['contact_info']
  address_info = contact_info['address']
  url_info = contact_info['url']
  park_guide.add_paragraph('Contact Info', 'Heading 1')
  park_guide.add_paragraph(address_info)
  park_guide.add_paragraph(url_info)

  #images 
  image_data = park_detail_response['nps_park_images']
  for image in image_data: 
    image_url = image['url']
    response = requests.get(image_url)
    image_content = response.content
    img = Image.open(BytesIO(image_content))
    park_guide.add_paragraph('Picture','Heading 1')
    park_guide.add_picture(BytesIO(image_content), width=docx.shared.Inches(4))
    print(image_url)

 
  #this is for activites it will look different since in the api it is bullets and multiple meaning?? We have to loop
  park_activities = park_detail_response['activities']
  park_guide.add_paragraph('Activities', 'Heading 1')
  for actvity in park_activities:
    park_guide.add_paragraph(actvity, style='List Bullet')



  park_guide.save('Testing.docx') #save doc file
